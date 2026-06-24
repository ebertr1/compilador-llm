#!/usr/bin/env python3
"""
Genera el documento .docx del proyecto APE10-11
Compilador en espanol con analisis lexico (AFD), sintactico, semantico e integracion LLM
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Global styles ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    hs.font.name = 'Calibri'

doc.styles['Heading 1'].font.size = Pt(18)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(12)

# ── Helper functions ───────────────────────────────────────────

def add_code(doc, code_text, caption=""):
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'f0f2f5')
    shading.set(qn('w:val'), 'clear')
    p.paragraph_format.element.get_or_add_pPr().append(shading)


def add_row(table, texts):
    row = table.add_row()
    for i, text in enumerate(texts):
        row.cells[i].text = text
        for p in row.cells[i].paragraphs:
            p.style = doc.styles['Normal']
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            for run in p.runs:
                run.font.size = Pt(9)


def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(2)


def add_trace(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)


# ══════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    "Diseno e Implementacion de un Compilador\n"
    "con Analisis Lexico (AFD), Sintactico y Semantico\n"
    "Integrando Modelos de Lenguaje (LLM)"
)
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run(
    "APE10-11 -- Lenguaje Tipado Estatico (ES)\n"
    "entero / flotante / cadena -- si/sino, mientras, para, imprimir, ENETRO\n"
    "Python + PLY + Flask + Gemma 3 1B (Ollama)"
)
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Compilador Web -- Backend\nJunio 2026")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. ABSTRACT
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Abstract (Resumen)", level=1)

doc.add_paragraph(
    "Este documento presenta el diseno e implementacion de un compilador completo para un lenguaje "
    "de programacion educativo con tipado estatico estricto en espanol que soporta tres tipos de datos: "
    "entero, flotante y cadena. El lenguaje incluye construcciones de control de flujo (si/sino, mientras, "
    "para), instrucciones de salida (imprimir) y definicion de tablas (ENETRO) con columnas tipadas. "
    "El compilador se estructura en tres fases fundamentales: (1) analisis lexico basado en un Automata "
    "Finito Determinista (AFD) que reconoce 24 tipos de tokens implementado con PLY, (2) analisis "
    "sintactico con gramatica libre de contexto y precedencia de operadores que construye un Arbol de "
    "Sintaxis Abstracta (AST) con 15 tipos de nodos, y (3) analisis semantico con tabla de simbolos "
    "y verificacion de tipos estricta sin coercion implicita, capaz de detectar 11 categorias de errores."
)
doc.add_paragraph(
    "Como contribucion novedosa, se integra el modelo de lenguaje Gemma 3 1B (ejecutado localmente via "
    "Ollama) para explicar los errores de compilacion en lenguaje natural tecnico en espanol. El sistema "
    "incluye una interfaz web de 4 paneles (Tokens, AST grafico, Compilador, IA) desarrollada con Flask, "
    "accesible desde cualquier dispositivo en la red local. El AST se renderiza como SVG directamente "
    "en el servidor sin dependencias externas, y tambien se genera codigo DOT como alternativa."
)
doc.add_paragraph(
    "Los resultados demuestran que el AFD procesa cadenas en tiempo lineal O(n), que la verificacion "
    "semantica detecta 11 categorias de errores (incluyendo validacion de columnas ENETRO y condiciones "
    "de tipo en estructuras de control), y que el LLM proporciona explicaciones consistentes y "
    "tecnicamente precisas con temperatura 0.1."
)

doc.add_heading("Palabras clave", level=3)
kw = "AFD, compilador, PLY, analisis semantico, Gemma 3, PLN, Python, JFLAP, tipado estatico, espanol, Ollama, Flask, ENETRO"
p = doc.add_paragraph()
run = p.add_run(kw)
run.bold = True
run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 2. INTRODUCCION
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Introduccion", level=1)

doc.add_paragraph(
    "La compilacion de lenguajes de programacion es un proceso que transforma codigo fuente escrito en un "
    "lenguaje de alto nivel a codigo maquina ejecutable. Este proceso se divide tradicionalmente en varias "
    "fases: analisis lexico, analisis sintactico, analisis semantico, generacion de codigo intermedio, "
    "optimizacion y generacion de codigo final."
)
doc.add_paragraph(
    "El analisis lexico constituye la primera fase y se encarga de convertir la secuencia de caracteres de "
    "entrada en una secuencia de tokens. Esta fase se implementa clasicamente mediante Automatas Finitos "
    "Deterministas (AFD), que procesan la entrada caracter por caracter en tiempo lineal y determinan a que "
    "categoria lexica pertenece cada secuencia (Hopcroft, Motwani y Ullman, 2006)."
)
doc.add_paragraph(
    "En anos recientes, los Modelos de Lenguaje de Gran Escala (LLMs) han demostrado capacidades notables "
    "en tareas de procesamiento de lenguaje natural, incluyendo la generacion y explicacion de codigo. "
    "Gemma, presentado por Google DeepMind (Team, 2024), ofrece modelos abiertos que pueden ejecutarse "
    "localmente, permitiendo su integracion en herramientas de desarrollo sin depender de servicios en la nube."
)
doc.add_paragraph(
    "Este proyecto integra ambas areas --la teoria clasica de compiladores con automatas y los modernos LLMs-- "
    "para crear un compilador en espanol que no solo detecte errores, sino que los explique en lenguaje natural "
    "tecnico. El compilador implementa un lenguaje con tipado estatico estricto que soporta tres tipos "
    "(entero, flotante, cadena) con verifacion completa de tipos en operaciones aritmeticas, comparaciones, "
    "asignaciones, y acceso a columnas de tablas ENETRO."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 3. PROBLEMATICA
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Problematica", level=1)

doc.add_heading("Contexto", level=2)
doc.add_paragraph(
    "En el diseno de un compilador para un lenguaje de programacion educativo con tipado estatico estricto "
    "en espanol, se necesita reconocer categorias fundamentales de tokens que constituyen los bloques lexicos "
    "del lenguaje. Ademas, el lenguaje debe soportar estructuras de control de flujo (si/sino, mientras, para), "
    "definicion de tablas con columnas tipadas (ENETRO), e instrucciones de salida (imprimir)."
)

doc.add_heading("El Lenguaje a Reconocer (Patrones Lexicos)", level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Token"
hdr[1].text = "Patron (Regex)"
hdr[2].text = "AFD -- Estados"
hdr[3].text = "Ejemplo"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

tokens_data = [
    ["Identificador", "[a-zA-Z_][a-zA-Z0-9_]*", "q0 a q_id_1 a q_id_2 (loop)", "x, _cont, var1"],
    ["Numero entero", r"\d+", "q0 a q_num_1 a q_num_2 (loop)", "42, 0, 123"],
    ["Numero flotante", r"\d+\.\d+", "q0 a q_num_1 a q_num_2 a q_float_1 a q_float_2", "3.14, 0.5"],
    ["String", r'"[^"]*"', "q0 a q_str_1 a q_str_2 a q_str_3", '"hola", "x"'],
    ["Operadores", r'[\+\-\*/%=;\(\)\{\},\.]', "q0 a q_op_1 (loop)", "+, -, *, /, {, }, ,"],
    ["Palabras clave ES", "entero|flotante|cadena|si|sino|mientras|para|imprimir|ENETRO", "Reserved en lexer", "si, mientras, ENETRO"],
]
for row_data in tokens_data:
    add_row(table, row_data)

doc.add_paragraph()

doc.add_heading("Justificacion del AFD como Solucion", level=2)
doc.add_paragraph(
    "Un Automata Finito Determinista (AFD) es la solucion correcta para el analisis lexico:"
)
add_bullet(doc, " Procesa cualquier cadena en tiempo lineal O(n) sin retroceso.", "Complejidad lineal: ")
add_bullet(doc, " No hay ambiguedad: cada estado y simbolo tiene un unico destino.", "Determinismo: ")
add_bullet(doc, " PLY construye internamente un AFD a partir de las expresiones regulares.", "Herramientas: ")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 4. OBJETIVOS
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Objetivos", level=1)

doc.add_heading("Objetivo General", level=2)
doc.add_paragraph(
    "Disenar e implementar un compilador completo en espanol con analisis lexico basado en un Automata "
    "Finito Determinista (AFD), analisis sintactico con gramatica libre de contexto y construccion de AST "
    "(15 tipos de nodos), y analisis semantico con verificacion de tipos estricta, integrando un modelo "
    "de lenguaje (Gemma 3 1B) mediante PLN para la explicacion de errores en espanol tecnico."
)

doc.add_heading("Objetivos Especificos", level=2)
objs = [
    "Definir formalmente la quintupla (Q, Sigma, delta, q0, F) del AFD lexico con 11 estados.",
    "Construir la tabla de transicion de estados del AFD lexico, verificando determinismo.",
    "Dibujar diagramas en JFLAP: AFN con epsilon-transiciones, AFD convertido y AFD minimizado.",
    "Demostrar el AFD con trazas de 5 cadenas aceptadas y 2 rechazadas.",
    "Implementar el lexer en Python con PLY para 24 tipos de tokens en espanol.",
    "Disenar gramatica BNF con precedencia de operadores e implementar parser con AST de 15 nodos.",
    "Implementar semantica con tabla de simbolos y 11 reglas de verificacion de tipos.",
    "Soportar estructuras de control: si/sino, mientras, para (con auto-declaracion de variable).",
    "Soportar tablas ENETRO con validacion de columnas y tipos.",
    "Integrar Gemma 3 1B (Ollama) con prompt engineering para explicaciones en espanol.",
    "Desarrollar interfaz web con 4 paneles (Tokens, AST, Compilador, IA) y renderizado SVG.",
]
for i, obj in enumerate(objs, 1):
    add_bullet(doc, obj)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 5. MARCO REFERENCIAL
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Marco Referencial", level=1)

doc.add_heading("Teoria de Automatas y Lenguajes Formales", level=2)
doc.add_paragraph(
    "Hopcroft, J. E., Motwani, R., & Ullman, J. D. (2006). Introduction to Automata Theory, Languages, "
    "and Computation (3rd ed.). Addison-Wesley."
)
doc.add_paragraph(
    "Referencia canonica para la teoria de automatas. Proporciona la definicion formal de los AFD como "
    "quintupla (Q, Sigma, delta, q0, F), los algoritmos de conversion AFN a AFD (construccion de "
    "subconjuntos), y el algoritmo de minimizacion de Hopcroft."
)
doc.add_paragraph(
    "Sipser, M. (2012). Introduction to the Theory of Computation (3rd ed.). Cengage Learning."
)
doc.add_paragraph(
    "Presenta la demostracion de que los lenguajes regulares son exactamente aquellos reconocidos por un AFD."
)

doc.add_heading("Diseno de Compiladores", level=2)
doc.add_paragraph(
    "Aho, A. V., Lam, M. S., Sethi, R., & Ullman, J. D. (2006). Compilers: Principles, Techniques, "
    "and Tools (2nd ed.). Addison-Wesley."
)
doc.add_paragraph(
    "Referencia clasica en diseno de compiladores. Define la arquitectura en fases que sigue este proyecto."
)
doc.add_paragraph(
    "Grune, D., & Jacobs, C. J. H. (2008). Parsering Techniques: A Practical Guide (2nd ed.). Springer."
)
doc.add_paragraph(
    "Tratamiento exhaustivo de tecnicas de analisis sintactico, incluyendo parsers LALR como PLY yacc."
)

doc.add_heading("Procesamiento de Lenguaje Natural con LLMs", level=2)
doc.add_paragraph(
    "Brown, T. B., et al. (2020). Language Models are Few-Shot Learners. NeurIPS, 33, 1877-1901."
)
doc.add_paragraph(
    "Presenta GPT-3 y demuestra que los LLMs pueden realizar tareas con few-shot learning. "
    "Base del prompt engineering utilizado en ollama_service.py."
)
doc.add_paragraph(
    "Team, G. (2024). Gemma: Open Models Based on Gemini Research and Technology. arXiv:2403.08295."
)
doc.add_paragraph(
    "Gemma es una familia de modelos abiertos de Google DeepMind. Gemma 3 1B tiene 1B parametros, "
    "se ejecuta localmente con Ollama, y proporciona generacion de texto de alta calidad."
)

doc.add_heading("Sistemas de Tipos Estaticos", level=2)
doc.add_paragraph(
    "Pierce, B. C. (2002). Types and Programming Languages. MIT Press."
)
doc.add_paragraph(
    "Referencia fundamental en sistemas de tipos. Define formalmente el tipado estatico y las reglas "
    "de verificacion que se implementan en semantic.py."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 6. METODOLOGIA
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Metodologia", level=1)

doc.add_paragraph(
    "La metodologia combina el enfoque clasico de diseno de automatas con desarrollo de software "
    "en cascada adaptada para compiladores."
)

doc.add_heading("Arquitectura del Proyecto", level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Fase"
hdr[1].text = "Actividad"
hdr[2].text = "Archivo(s)"
hdr[3].text = "Salida"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)

met_data = [
    ["1. Lexico", "AFD con PLY, 24 tokens, palabras clave ES", "lexer.py", "Lista de tokens"],
    ["2. Sintactico", "Gramatica BNF, LALR, precedencia", "parser.py, ast_nodes.py", "AST (15 tipos)"],
    ["3. Semantico", "Tabla de simbolos, 11 reglas de tipos", "semantic.py, symbols.py", "Errores / exito"],
    ["4. LLM", "Prompt engineering, Ollama, gemma3:1b", "ollama_service.py", "Explicacion en ES"],
    ["5. Web", "Flask, 4 paneles, SVG del AST", "web_api.py, index.html", "Interfaz web"],
    ["6. Visualizacion", "SVG servidor, DOT fallback", "ast_to_svg.py, ast_to_dot.py", "Arbol AST grafico"],
]
for row_data in met_data:
    add_row(table, row_data)

doc.add_paragraph()

doc.add_heading("Fase 1: Analisis Lexico", level=2)
doc.add_paragraph(
    "Se implemento con PLY (Python Lex-Yacc). Cada regla t_XXX define una expresion regular que PLY "
    "convierte internamente en un AFD. Las palabras clave reservadas (entero, flotante, cadena, si, "
    "sino, mientras, para, imprimir, ENETRO) se reconocen con prioridad sobre los identificadores."
)

with open(os.path.join(os.path.dirname(__file__), 'lexer.py')) as f:
    lexer_code = f.read()
add_code(doc, lexer_code, "Codigo completo del lexer (lexer.py):")

doc.add_paragraph()
doc.add_paragraph("Cada regla del lexer corresponde a una transicion del AFD:")
add_bullet(doc, " t_FLOAT_NUMBER: reconoce flotantes y convierte a float.", "")
add_bullet(doc, " t_NUMBER: reconoce enteros y convierte a int.", "")
add_bullet(doc, " t_TEXT: reconoce strings \"...\" y extrae el contenido.", "")
add_bullet(doc, " t_ID: reconoce identificadores; reserved distingue palabras clave.", "")
add_bullet(doc, " t_LE, t_GE, t_EQ, t_NE: operadores de dos caracteres (mayor prioridad).", "")
add_bullet(doc, " t_LBRACE, t_RBRACE, t_COMMA, t_DOT: nuevos tokens para bloques y ENETRO.", "")
add_bullet(doc, " t_comment_hash, t_comment_slash: ignoran comentarios # y //.", "")

doc.add_heading("Fase 2: Analisis Sintactico", level=2)
doc.add_paragraph("Gramatica libre de contexto con producciones:")
add_code(doc, """program      → program statement | statement | ε
statement    → var_decl | assign | if_stmt | while_stmt
             | for_stmt | print_stmt | enetro_def
var_decl     → type ID = expression ;
             | type ID ;
assign       → ID = expression ;
if_stmt      → SI LPAREN expression RPAREN block (SINO block)?
while_stmt   → MIENTRAS LPAREN expression RPAREN block
for_stmt     → PARA LPAREN assign? ; expression ; assign? RPAREN block
print_stmt   → IMPRIMIR LPAREN expression RPAREN ;
enetro_def   → ENETRO ID LBRACE enetro_fields RBRACE
enetro_field → type ID = expression ;
block        → LBRACE statements RBRACE
type         → ENTERO | FLOTANTE | CADENA""")

doc.add_paragraph()
doc.add_paragraph("Precedencia de operadores (de menor a mayor):")
add_code(doc, """==  !=
<   >  <=  >=
+   -
*   /   %
- (unario)     (mayor)""")

doc.add_paragraph()

with open(os.path.join(os.path.dirname(__file__), 'parser.py')) as f:
    parser_code = f.read()
add_code(doc, parser_code, "Codigo completo del parser (parser.py):")

doc.add_paragraph()
doc.add_paragraph("El parser utiliza PLY yacc con las siguientes caracteristicas:")
add_bullet(doc, " 15 tipos de nodos AST: Program, Block, VarDecl, Assign, BinOp, UnaryOp, Literal, VarRef, IfStmt, WhileStmt, ForStmt, PrintStmt, EnetroDef, EnetroField, EnetroAccess.", "")
add_bullet(doc, " Los errores sintacticos se acumulan en syntax_errors sin detener el analisis.", "")
add_bullet(doc, " Manejo de bloques { } para estructuras de control.", "")
add_bullet(doc, " Soporte de ENETRO con campos tipados y valores iniciales.", "")

doc.add_heading("Fase 3: Analisis Semantico", level=2)

with open(os.path.join(os.path.dirname(__file__), 'ast_nodes.py')) as f:
    ast_code = f.read()
add_code(doc, ast_code, "Nodos del AST (ast_nodes.py):")

doc.add_paragraph()

with open(os.path.join(os.path.dirname(__file__), 'compiler_error.py')) as f:
    cerr = f.read()
add_code(doc, cerr, "Estructura de error (compiler_error.py):")

doc.add_paragraph()

with open(os.path.join(os.path.dirname(__file__), 'symbols.py')) as f:
    sym = f.read()
add_code(doc, sym, "Tabla de simbolos (symbols.py):")

doc.add_paragraph()

with open(os.path.join(os.path.dirname(__file__), 'semantic.py')) as f:
    sem = f.read()
add_code(doc, sem, "Analizador semantico (semantic.py):")

doc.add_paragraph()
doc.add_paragraph("El analizador semantico implementa las siguientes 11 reglas de verificacion:")

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Regla"
hdr[1].text = "Condicion"
hdr[2].text = "Ejemplo de error"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)

rules = [
    ["No redeclaracion", "VarDecl: name ya existe en tabla", "'x' ya declarada"],
    ["Variable existe", "Assign: name no existe", "'no_existe' no declarada"],
    ["Tipo compatible (decl)", "expr_type == var_type", "'x' es entero pero expresion es cadena"],
    ["Tipo compatible (asig)", "expr_type == var_type", "tipo incompatible en asignacion"],
    ["[+] tipos iguales", "str+str, int+int, float+float", "'+' invalida entre entero y flotante"],
    ["[%] solo entero", "int % int", "'%' solo valido para entero"],
    ["Comparacion tipos iguales", "mismo tipo ambos lados", "no se puede comparar X con Y"],
    ["-, *, / numericos", "ambos del mismo tipo numerico", "operacion invalida entre tipos"],
    ["Unario solo numerico", "operando es entero o flotante", "no se puede aplicar '-' a cadena"],
    ["Condicion entera", "si/mientras requiren entero", "condicion del 'si' debe ser entero"],
    ["Columna ENETRO existe", "acceso a columna inexistente", "'edad' no existe en ENETRO 'usuarios'"],
]
for r in rules:
    add_row(table, r)

doc.add_paragraph()
doc.add_paragraph("Ademas, el 'para' auto-declara la variable de iteracion como entero si no existe previamente.")

doc.add_heading("Fase 4: Integracion LLM", level=2)

with open(os.path.join(os.path.dirname(__file__), 'ollama_service.py')) as f:
    ollama = f.read()
add_code(doc, ollama, "Servicio de explicacion con IA (ollama_service.py):")

doc.add_paragraph()
doc.add_paragraph("Prompt engineering en dos modos:")

doc.add_heading("Modo errores", level=3)
add_code(doc, """Eres un ingeniero senior en compiladores.
Responde en espanol, sin asteriscos, sin caracteres de markdown.
Recibes una lista de errores reales detectados por el compilador.
Analiza UNICAMENTE los errores listados, sin inventar otros.

Por cada error de la lista, responde con:
  Tipo: SemanticError o SyntaxError
  Linea: X
  Causa: explicacion tecnica concisa
  Regla violada: nombre de la regla semantica violada""")

doc.add_paragraph()
doc.add_heading("Modo exito", level=3)
add_code(doc, """Responde en espanol, sin asteriscos, sin caracteres de markdown.
El codigo ha pasado compilacion exitosamente.
Responde con:
  Estado: Compilacion exitosa
  Tabla de simbolos: variable: tipo, variable: tipo
  Verificacion de tipos: APROBADA
  Estructuras: descripcion breve de ENETRO, condicionales, bucles usados""")

doc.add_paragraph()
doc.add_paragraph("Parametros de conexion a Ollama:")
add_bullet(doc, " Modelo: gemma3:1b (1B parametros, ejecucion local)", "")
add_bullet(doc, " Temperature: 0.1 (baja creatividad, respuestas deterministas)", "")
add_bullet(doc, " num_predict: 500 (maximo de tokens de salida)", "")
add_bullet(doc, " Timeout: (3s connect, 60s read)", "")

doc.add_heading("Fase 5: Visualizacion del AST", level=2)

with open(os.path.join(os.path.dirname(__file__), 'ast_to_svg.py')) as f:
    svg_code = f.read()
add_code(doc, svg_code, "Renderizador SVG del AST (ast_to_svg.py):")

doc.add_paragraph()
doc.add_paragraph(
    "El modulo ast_to_svg.py genera un SVG del arbol AST directamente desde Python sin dependencias externas. "
    "Usa un layout recursivo que posiciona los nodos en el plano (x, y) y dibuja cajas rectangulares "
    "con colores por tipo de nodo, sombras suaves, y flechas curvas entre padre e hijos. "
    "El modulo ast_to_dot.py genera codigo DOT como alternativa para herramientas externas."
)

doc.add_heading("Fase 6: Frontend Web", level=2)

with open(os.path.join(os.path.dirname(__file__), 'web_api.py')) as f:
    web_code = f.read()
add_code(doc, web_code, "Servidor web Flask (web_api.py):")

doc.add_paragraph()

with open(os.path.join(os.path.dirname(__file__), 'main.py')) as f:
    main_code = f.read()
add_code(doc, main_code, "Pipeline del compilador (main.py):")

doc.add_paragraph()
doc.add_paragraph("El pipeline del compilador sigue este flujo:")
add_code(doc, """compile_code(code):
  1. get_tokens(code)        -- tokens lista
  2. parse(code)             -- AST
  3. if syntax_errors:       -- explain_with_llm(errors)
  4. analyze(ast)            -- semantica
  5. if semantic errors:     -- explain_with_llm(errors)
  6. success                 -- explain_with_llm(empty)
  7. return { tokens, ast, dot, svg, errors, llm_explanation }""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 7. RESULTADOS
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Resultados", level=1)

doc.add_heading("Definicion Formal del AFD (Quintupla)", level=2)
doc.add_paragraph("El AFD disenado para el analisis lexico se define formalmente como:")
add_code(doc, """AFD_Lex = (Q, Sigma, delta, q0, F)

Q = { q0, q_id_1, q_id_2, q_num_1, q_num_2, q_float_1, q_float_2,
      q_str_1, q_str_2, q_str_3, q_op_1 }

Sigma = { a-z, A-Z, 0-9, _, ., ", +, -, *, /, %, =, ;, (, ), {, }, ,,
          #, /, <, >, !, espacio }

q0 = q0

F = { q_id_2, q_num_2, q_float_2, q_str_3, q_op_1 }""")

doc.add_paragraph()
doc.add_heading("Descripcion de Estados", level=3)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Estado"
hdr[1].text = "Nombre"
hdr[2].text = "Tipo"
hdr[3].text = "Descripcion"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)

state_data = [
    ["q0", "Inicial", "Inicial", "Distribuye a sub-automatas via epsilon"],
    ["q_id_1", "ID inicio", "Intermedio", "Recibe primera letra/guion bajo"],
    ["q_id_2", "ID final", "Aceptacion", "Acepta identificadores completos"],
    ["q_num_1", "Numero inicio", "Intermedio", "Recibe primer digito"],
    ["q_num_2", "Numero final", "Aceptacion", "Acepta enteros; conecta a flotante via '.'"],
    ["q_float_1", "Flotante decimal", "Intermedio", "Recibe el punto decimal"],
    ["q_float_2", "Flotante final", "Aceptacion", "Acepta flotantes completos"],
    ["q_str_1", "String inicio", "Intermedio", "Espera comilla de apertura"],
    ["q_str_2", "String contenido", "Intermedio", "Acumula caracteres del string"],
    ["q_str_3", "String final", "Aceptacion", "Acepta string con comilla de cierre"],
    ["q_op_1", "Operador", "Aceptacion", "Acepta cualquier operador/simbolo"],
]
for r in state_data:
    add_row(table, r)

doc.add_paragraph()

doc.add_heading("Tabla de Transicion delta", level=2)

table = doc.add_table(rows=12, cols=7)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Estado \\ Simbolo", "letra/_", "digito", ".", '"', "operador", "otro"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(8)

trans_data = [
    ["q0", "q_id_1", "q_num_1", "--", "q_str_1", "q_op_1", "--"],
    ["q_id_1", "q_id_2", "--", "--", "--", "--", "--"],
    ["q_id_2", "q_id_2", "q_id_2", "--", "--", "--", "--"],
    ["q_num_1", "--", "q_num_2", "--", "--", "--", "--"],
    ["q_num_2", "--", "q_num_2", "q_float_1", "--", "--", "--"],
    ["q_float_1", "--", "q_float_2", "--", "--", "--", "--"],
    ["q_float_2", "--", "q_float_2", "--", "--", "--", "--"],
    ["q_str_1", "--", "--", "--", "q_str_2", "--", "--"],
    ["q_str_2", "q_str_2", "q_str_2", "q_str_2", "q_str_3", "q_str_2", "q_str_2"],
    ["q_str_3", "--", "--", "--", "--", "--", "--"],
    ["q_op_1", "--", "--", "--", "--", "q_op_1", "--"],
]
for ri, rd in enumerate(trans_data, start=1):
    for ci, val in enumerate(rd):
        cell = table.rows[ri].cells[ci]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)

doc.add_paragraph()

doc.add_heading("Diagrama de Estados (JFLAP)", level=2)
doc.add_paragraph(
    "Los diagramas de estados generados con JFLAP se encuentran en el directorio jflap/: "
    "AFN.jff (Automata Finito No-Determinista con epsilon-transiciones), "
    "AFD.jff (Automata Finito Determinista convertido), "
    "y MIN.jff (Automata Finito Determinista Minimizado con Hopcroft)."
)

doc.add_page_break()

doc.add_heading("Validacion con Trazas", level=2)
doc.add_paragraph(
    "Trazas de ejecucion del AFD para 5 cadenas aceptadas y 2 rechazadas:"
)

doc.add_heading("Cadena 1: 'x' (identificador)", level=3)
add_trace(doc, [
    "(q0, epsilon) a q_id_1",
    "(q_id_1, 'x') a q_id_2  <- FINAL OK",
    "Token: ID('x')"
])

doc.add_heading("Cadena 2: '123' (entero)", level=3)
add_trace(doc, [
    "(q0, epsilon) a q_num_1",
    "(q_num_1, '1') a q_num_2",
    "(q_num_2, '2') a q_num_2",
    "(q_num_2, '3') a q_num_2  <- FINAL OK",
    "Token: NUMBER(123)"
])

doc.add_heading("Cadena 3: '3.14' (flotante)", level=3)
add_trace(doc, [
    "(q0, epsilon) a q_num_1",
    "(q_num_1, '3') a q_num_2",
    "(q_num_2, '.') a q_float_1",
    "(q_float_1, '1') a q_float_2",
    "(q_float_2, '4') a q_float_2  <- FINAL OK",
    "Token: FLOAT_NUMBER(3.14)"
])

doc.add_heading("Cadena 4: '\"hola\"' (string)", level=3)
add_trace(doc, [
    "(q0, epsilon) a q_str_1",
    "(q_str_1, '\"') a q_str_2",
    "(q_str_2, 'h') a q_str_2",
    "(q_str_2, 'o') a q_str_2",
    "(q_str_2, 'l') a q_str_2",
    "(q_str_2, 'a') a q_str_2",
    "(q_str_2, '\"') a q_str_3  <- FINAL OK",
    "Token: TEXT('hola')"
])

doc.add_heading("Cadena 5: '+' (operador)", level=3)
add_trace(doc, [
    "(q0, epsilon) a q_op_1",
    "(q_op_1, '+') a q_op_1  <- FINAL OK",
    "Token: PLUS"
])

doc.add_heading("Cadena Rechazada 1: '1x' (numero + letra)", level=3)
add_trace(doc, [
    "(q0, epsilon) a q_num_1",
    "(q_num_1, '1') a q_num_2",
    "(q_num_2, 'x') a SIN TRANSICION",
    "Resultado: error lexico"
])

doc.add_heading("Cadena Rechazada 2: '\"abc' (string sin cerrar)", level=3)
add_trace(doc, [
    "(q0, epsilon) a q_str_1",
    "(q_str_1, '\"') a q_str_2",
    "(q_str_2, 'a') a q_str_2",
    "(q_str_2, 'b') a q_str_2",
    "(q_str_2, 'c') a q_str_2",
    "(fin entrada) a NO ESTA EN F",
    "Resultado: error lexico -- string sin cerrar"
])

doc.add_page_break()

doc.add_heading("Resultados del Compilador Completo", level=2)

doc.add_heading("Ejemplo: Codigo Correcto", level=3)
with open(os.path.join(os.path.dirname(__file__), '..', 'examples', 'correct_sp.txt')) as f:
    correct_code = f.read()
add_code(doc, correct_code, "Codigo correcto (correct_sp.txt) -- 10 constructos:")

doc.add_paragraph()
doc.add_paragraph("Al compilar, el resultado es exitoso con 0 errores. Se genera:")
add_bullet(doc, " Lista de tokens con tipos y valores", "")
add_bullet(doc, " SVG del arbol AST con nodos coloreados por categoria", "")
add_bullet(doc, " Tabla de simbolos con todas las variables y tipos", "")
add_bullet(doc, " Explicacion de la IA en espanol confirmando compilacion exitosa", "")

doc.add_paragraph()
doc.add_paragraph("El SVG del AST renderiza nodos con colores profesionales:")
add_bullet(doc, " Program / Block: azul oscuro #2C3E50", "")
add_bullet(doc, " VarDecl / Assign: verde #27AE60", "")
add_bullet(doc, " BinOp / UnaryOp: rojo #E74C3C", "")
add_bullet(doc, " IfStmt / WhileStmt / ForStmt: naranja #E67E22", "")
add_bullet(doc, " PrintStmt: purpura #8E44AD", "")
add_bullet(doc, " EnetroDef / EnetroField: azul #2980B9", "")
add_bullet(doc, " Literal / VarRef: gris #7F8C8D", "")

doc.add_paragraph()

doc.add_heading("Ejemplo: Codigo con Errores", level=3)
with open(os.path.join(os.path.dirname(__file__), '..', 'examples', 'errors_sp.txt')) as f:
    errors_code = f.read()
add_code(doc, errors_code, "Codigo con errores (errors_sp.txt) -- 11 errores semanticos:")

doc.add_paragraph()
doc.add_paragraph("El compilador detecta 11 errores semanticos:")

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Linea"
hdr[1].text = "Error"
hdr[2].text = "Regla violada"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

err_data = [
    ["5", "tipo incompatible: 'y' es entero pero expresion es cadena", "Type mismatch"],
    ["8", "variable 'no_existe' no declarada", "Undeclared identifier"],
    ["12", "variable 'pi' ya declarada", "Redeclaration"],
    ["17", "operacion '+' invalida entre entero y flotante", "Mixed types binop"],
    ["20", "'%' solo valido para entero, no entre flotante y entero", "Modulo non-int"],
    ["24", "no se puede aplicar '-' a tipo cadena", "Unary non-numeric"],
    ["28", "operacion '+' invalida entre cadena y entero", "Mixed types concat"],
    ["34", "columna 'apellido' no existe en ENETRO 'usuarios'", "ENETRO column"],
    ["38", "tipo incompatible en columna 'codigo' de ENETRO 'productos'", "ENETRO type"],
    ["42", "condicion del 'si' debe ser entero, no 'cadena'", "If condition type"],
    ["47", "condicion del 'mientras' debe ser entero, no 'cadena'", "While condition type"],
]
for r in err_data:
    add_row(table, r)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 8. DISCUSION
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Discusion", level=1)

doc.add_heading("AFD vs Expresiones Regulares", level=2)
doc.add_paragraph(
    "PLY convierte internamente las expresiones regulares a un AFD, combinando la familiaridad de las "
    "regex con la eficiencia O(n) del automata. La ventaja de PLY es que abstrae la construccion manual "
    "del AFD, pero el fundamento teorico sigue siendo el mismo."
)

doc.add_heading("Minimizacion de Estados", level=2)
doc.add_paragraph(
    "El AFN original (11 estados + epsilon-transiciones) se convirtio a AFD y se minimizo con Hopcroft "
    "en JFLAP, eliminando estados inalcanzables y equivalentes para optimizar memoria."
)

doc.add_heading("Efectividad del LLM (Gemma 3 1B)", level=2)
doc.add_paragraph("La integracion del LLM presenta los siguientes hallazgos:")
add_bullet(doc, " Temperature 0.1 produce respuestas consistentes y deterministas.", "Baja temperatura: ")
add_bullet(doc, " El prompt estructurado en espanol evita contenido no solicitado.", "Formato: ")
add_bullet(doc, " 500 tokens son suficientes para explicar multiples errores.", "Tamano: ")
add_bullet(doc, " El LLM nunca participa en decisiones de compilacion.", "Seguridad: ")

doc.add_heading("Lenguaje en Espanol", level=2)
doc.add_paragraph(
    "La decision de usar palabras clave en espanol (entero, flotante, cadena, si, sino, mientras, para, "
    "imprimir, ENETRO) hace que el lenguaje sea accesible para hispanohablantes que se inician en la "
    "programacion. El compilador mantiene el rigor del tipado estatico estricto sin sacrificar claridad."
)

doc.add_heading("Renderizado SVG del AST", level=2)
doc.add_paragraph(
    "El modulo ast_to_svg.py genera el arbol AST como SVG directamente en el servidor sin requerir "
    "Graphviz, CDN ni conexion a internet. Esto asegura que el arbol siempre sea visible, incluso en "
    "entornos sin acceso a redes externas. Los nodos se colorean por categoria semantica y se conectan "
    "con flechas curvas para una visualizacion clara y profesional."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 9. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Conclusiones", level=1)

doc.add_paragraph(
    "1. El AFD disenado reconoce exactamente los 5 tipos de tokens del lenguaje (identificadores, "
    "enteros, flotantes, strings y operadores) sin ambiguedad. Las trazas confirman que las cadenas "
    "validas son correctamente aceptadas y las invalidas rechazadas."
)
doc.add_paragraph(
    "2. El compilador procesa cualquier programa en tiempo lineal O(n) respecto a la longitud del "
    "codigo de entrada."
)
doc.add_paragraph(
    "3. El automata minimizado (MIN.jff) tiene el numero optimo de estados tras aplicar Hopcroft."
)
doc.add_paragraph(
    "4. La integracion de Gemma 3 1B proporciona explicaciones tecnicas consistentes de los errores "
    "en espanol, mejorando la experiencia del desarrollador."
)
doc.add_paragraph(
    "5. El sistema de tipos estaticos estrictos sin coercion implicita detecta 11 categorias de errores "
    "semanticos: redeclaracion, variable no declarada, tipos incompatibles en asignacion, tipos mixtos "
    "en operaciones binarias, modulo no entero, operador unario en no numerico, comparacion entre tipos "
    "distintos, concatenacion invalida, condicion no entera en si/mientras, columna ENETRO inexistente, "
    "y tipo incompatible en columna ENETRO."
)
doc.add_paragraph(
    "6. El renderizado SVG del AST en el servidor garantiza que el arbol sea visible sin depender de "
    "conexion a internet ni librerias externas."
)
doc.add_paragraph(
    "7. La interfaz web de 4 paneles (Tokens, AST, Compilador, IA) proporciona una experiencia "
    "completa de compilacion accesible desde cualquier navegador en la red local."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 10. RECOMENDACIONES
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Recomendaciones", level=1)

doc.add_heading("Extensiones del Lenguaje", level=2)
doc.add_paragraph("Se recomienda extender el lenguaje con:")
add_bullet(doc, " Funciones con parametros y retorno", "")
add_bullet(doc, " Arreglos y estructuras de datos compuestas", "")
add_bullet(doc, " Iteracion sobre ENETRO (foreach)", "")
add_bullet(doc, " Operadores logicos: &&, ||, !", "")

doc.add_heading("Mejora del LLM", level=2)
add_bullet(doc, " Migrar a un modelo mas grande (Gemma 3 7B o 27B) para explicaciones mas detalladas.", "")
add_bullet(doc, " Implementar fine-tuning con datos especificos de errores del compilador.", "")
add_bullet(doc, " Cachear respuestas del LLM para errores identicos.", "")

doc.add_heading("Visualizacion", level=2)
add_bullet(doc, " Agregar animacion de trazas del AFD en la interfaz web.", "")
add_bullet(doc, " Implementar REPL interactivo para probar expresiones.", "")
add_bullet(doc, " Mostrar tabla de transicion del AFD interactiva.", "")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 11. REFERENCIAS
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Referencias", level=1)

refs = [
    "Aho, A. V., Lam, M. S., Sethi, R., & Ullman, J. D. (2006). Compilers: Principles, Techniques, and Tools (2nd ed.). Addison-Wesley.",
    "Brown, T. B., Mann, B., Ryder, N., et al. (2020). Language Models are Few-Shot Learners. Advances in Neural Information Processing Systems, 33, 1877-1901.",
    "Grune, D., & Jacobs, C. J. H. (2008). Parsing Techniques: A Practical Guide (2nd ed.). Springer.",
    "Hopcroft, J. E., Motwani, R., & Ullman, J. D. (2006). Introduction to Automata Theory, Languages, and Computation (3rd ed.). Addison-Wesley.",
    "Pierce, B. C. (2002). Types and Programming Languages. MIT Press.",
    "Sipser, M. (2012). Introduction to the Theory of Computation (3rd ed.). Cengage Learning.",
    "Team, G. (2024). Gemma: Open Models Based on Gemini Research and Technology. arXiv preprint arXiv:2403.08295.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"[{i}] {ref}")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 12. ANEXOS
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Anexos", level=1)

doc.add_heading("Anexo A: Diagramas JFLAP", level=2)
doc.add_paragraph("Directorios jflap/ contiene:")
add_bullet(doc, " jflap/AFN.jff -- AFN con epsilon-transiciones", "")
add_bullet(doc, " jflap/AFD.jff -- AFD convertido", "")
add_bullet(doc, " jflap/MIN.jff -- AFD minimizado", "")

doc.add_heading("Anexo B: Codigo Fuente", level=2)
doc.add_paragraph("backend/ contiene los modulos del compilador:")
add_bullet(doc, " lexer.py -- Analizador lexico (24 tokens, palabras clave ES)", "")
add_bullet(doc, " parser.py -- Analizador sintactico (gramatica LALR)", "")
add_bullet(doc, " semantic.py -- Analizador semantico (11 reglas)", "")
add_bullet(doc, " symbols.py -- Tabla de simbolos", "")
add_bullet(doc, " ast_nodes.py -- 15 clases de nodos AST", "")
add_bullet(doc, " compiler_error.py -- Estructura de error", "")
add_bullet(doc, " ollama_service.py -- Integracion Gemma 3 1B", "")
add_bullet(doc, " main.py -- Pipeline del compilador", "")
add_bullet(doc, " web_api.py -- Servidor Flask", "")
add_bullet(doc, " ast_to_svg.py -- Renderizado SVG del AST", "")
add_bullet(doc, " ast_to_dot.py -- Generacion DOT (fallback)", "")
add_bullet(doc, " templates/index.html -- Frontend 4 paneles", "")
add_bullet(doc, " run_server.sh -- Script de inicio", "")

doc.add_heading("Anexo C: Ejemplos de Codigo", level=2)
add_bullet(doc, " examples/correct_sp.txt -- Codigo correcto con 10 constructos", "")
add_bullet(doc, " examples/errors_sp.txt -- Codigo con 11 errores semanticos", "")

doc.add_heading("Anexo D: Instrucciones de Ejecucion", level=2)
doc.add_paragraph("Para ejecutar el compilador en modo servidor web:")
add_code(doc, """# 1. Instalar dependencias
pip install -r backend/requirements.txt

# 2. Iniciar Ollama (opcional, para IA)
ollama pull gemma3:1b

# 3. Ejecutar servidor
cd backend
bash run_server.sh

# 4. Abrir navegador en http://localhost:8080""")

doc.add_paragraph()
doc.add_paragraph("Para interfaz de linea de comandos:")
add_code(doc, """cd backend
# Modo interactivo
python3 main.py

# Modo JSON
echo 'entero x = 5;' | python3 main.py --json""")

# ── Save ───────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), '..', 'Documento_Compilador_APE10-11.docx')
doc.save(output_path)
print(f"OK Documento generado: {output_path}")
